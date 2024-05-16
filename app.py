from docx import Document
import os
from thefuzz import fuzz
import streamlit as st
import time
import subprocess
import string
from datetime import datetime

SUPPORTED_FILES = {".docx"}


class Paragraph:
    score = 0

    def __init__(self, text, file_path, last_modified) -> None:
        self.text = text
        self.file_path = file_path
        self.last_modified = last_modified

    def __repr__(self) -> str:
        date = datetime.fromtimestamp(self.last_modified).date().strftime("%Y-%m-%d")
        return f"""\nFile path: {self.file_path}\nText: {self.text}\nLast Modified: {date}\n"""


class Indexer:

    paragraphs = []

    def __init__(self, top_dir) -> None:
        self.top_dir = top_dir

    def index(self):
        print(f"Start indexing directory {self.top_dir}...")
        start = time.time()
        for (root, _, files) in os.walk(self.top_dir, topdown=True):
            for file in files:
                if os.path.splitext(file)[1] == ".docx" and not file.startswith("~$"):
                    full_path = os.path.join(root, file)
                    doc = Document(full_path)
                    last_modified = os.path.getmtime(full_path)
                    for p in doc.paragraphs:
                        words = p.text.split()
                        num_pun = count_chars(p.text, string.punctuation)
                        if not num_pun:
                            num_pun = 1
                        # only read paragraphs with more than five words and have a word/punctuation ratio of 3
                        # todo enable customization of these values
                        if len(words) > 5 and len(words)/num_pun > 3:
                            self.paragraphs.append(
                                Paragraph(p.text, full_path, last_modified))
                # todo: add pdf support
                # if os.path.splitext(file)[1] == ".pdf":
                #     full_path = os.path.join(root,file)
                #     doc = Document(full_path)
                #     for p in doc.paragraphs:
                #         if len(p.text.split())>5:
                #             self.paragraphs.append(Paragraph(p.text, full_path))
        print(f"Completed indexing in {time.time() - start} seconds!")

    def search(self, query, num_results=3):
        ratios = []
        results = []
        for p in self.paragraphs:
            score = fuzz.token_set_ratio(p.text, query)
            ratios.append(score)
            p.score = score
        sorted_indexes = sorted(
            range(len(ratios)), key=lambda k: ratios[k], reverse=True)

        for i in sorted_indexes[:num_results]:
            results.append(self.paragraphs[i])
        return results

def count_chars(string: str, chars: set):
    """helper function to count the number times the string contains
      a character in the set """
    count = 0
    for c in string:
        if c in chars:
            count +=1
    return count


@st.cache_resource(show_spinner=False)
def init_indexer(dir_path):
    indexer = Indexer(dir_path)
    indexer.index()
    return indexer


if __name__ == "__main__":
    directory = os.environ.get("TOP_DIR")
    st.title("Document Search")
    search_input = st.sidebar.text_input("Enter keywords...")
    num_results = st.sidebar.selectbox("Number of results", [10, 20, 50, 100])
    if not directory:
        raise RuntimeError("Environment variable `TOP_DIR` must be set!")
    with st.spinner(f'Indexing directory {directory}...'):
        start = time.time()
        indexer = init_indexer(directory)
    st.success(f'Completed Indexing in {int(time.time() - start)} seconds')
    if search_input:
        results = indexer.search(search_input, num_results=num_results)
        for i, paragraph in enumerate(results):
            date = datetime.fromtimestamp(paragraph.last_modified).date().strftime("%Y-%m-%d")
            text = f"""
            **File**: {paragraph.file_path}

            **Line**: {paragraph.text}

            **Last Modified**: {date}

            **Match Score**: {paragraph.score}
            """
            st.markdown(text)
            st.divider()
